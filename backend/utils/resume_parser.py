# 文件路径: D:\兰花花\talentpool\backend\utils\resume_parser.py

import re
from typing import List, Dict, Optional
import PyPDF2
import docx
from docx.text.paragraph import Paragraph

def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    从 PDF 文件中提取所有文本内容。

    :param file_path: PDF 文件的路径。
    :return: 提取的文本字符串，如果失败则返回 None。
    """
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return None

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    从 DOCX 文件中提取所有文本内容。

    :param file_path: DOCX 文件的路径。
    :return: 提取的文本字符串，如果失败则返回 None。
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        # 处理段落
        for para in doc.paragraphs:
            full_text.append(para.text)
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text).strip()
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return None

def parse_resume_content(text: str) -> Dict[str, any]:
    """
    从简历文本中解析结构化信息。

    注意：这是一个基于正则表达式的简化解析器。对于复杂或非标准的简历，
    效果可能不佳。在生产环境中，通常会使用更高级的 NLP 模型或专门为
    简历解析设计的库。

    :param text: 从简历文件中提取的原始文本。
    :return: 包含解析出的个人信息的字典。
    """
    if not text:
        return {}

    # 使用字典来存储解析结果，避免返回None值
    result = {
        "name": None,
        "email": [],
        "phone": [],
        "skills": [],
        "experience": [],  # 尝试提取工作经验
        "education": []    # 尝试提取教育背景
    }

    # 1. 提取姓名 (通常在第一行或特定标题附近)
    # 假设姓名是第一行非空的短行，并且不是常见的标题或联系方式
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        # 简单的启发式规则：第一行，并且不是"简历"等词
        potential_name = lines[0]
        if len(potential_name.split()) <= 4 and not re.search(r'(resume|个人简历|姓名|Name)', potential_name, re.IGNORECASE):
            result["name"] = potential_name

    # 2. 提取邮箱
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    result["email"] = list(set(re.findall(email_pattern, text)))

    # 3. 提取电话号码 (匹配中国大陆11位手机号)
    phone_pattern = r'1[3-9]\d{9}'
    result["phone"] = list(set(re.findall(phone_pattern, text)))

    # 4. 提取技能 (这是一个非常简化的示例，实际应用中需要一个技能词典)
    # 假设简历中有一个 "Skills:" 或 "技能:" 或 "技术栈" 标题
    # 然后提取该行之后的内容
    skills_section_pattern = re.compile(r'(skills|技能|技术栈|技术能力):?\s*([\s\S]*?)(?=\n\n|\Z)', re.IGNORECASE)
    skills_match = skills_section_pattern.search(text)
    if skills_match:
        # 获取技能部分的文本
        skills_text = skills_match.group(2)
        # 按逗号、空格或换行符分割，并清理
        skills = re.split(r'[,，\s]+', skills_text.strip())
        # 过滤掉空字符串和过短的词
        result["skills"] = [s for s in skills if s and len(s) > 1]

    # 5. 提取工作经验 (示例：查找 "工作经历" 或 "Experience" 后面的内容)
    experience_pattern = re.compile(r'(工作经历|Experience):?\s*([\s\S]*?)(?=\n\n(工作经历|Experience|教育背景|Education|Skills|技能)|\Z)', re.IGNORECASE)
    experience_match = experience_pattern.search(text)
    if experience_match:
        # 简单地按换行符分割，作为每段经历
        exp_text = experience_match.group(2).strip()
        result["experience"] = [line.strip() for line in exp_text.split('\n') if line.strip()]

    # 6. 提取教育背景 (示例：查找 "教育背景" 或 "Education" 后面的内容)
    education_pattern = re.compile(r'(教育背景|Education):?\s*([\s\S]*?)(?=\n\n(工作经历|Experience|教育背景|Education|Skills|技能)|\Z)', re.IGNORECASE)
    education_match = education_pattern.search(text)
    if education_match:
        edu_text = education_match.group(2).strip()
        result["education"] = [line.strip() for line in edu_text.split('\n') if line.strip()]

    return result

